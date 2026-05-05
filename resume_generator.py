import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def convert_profile_to_resume_format(profile):
    """Convert app profile schema into ATS-resume schema."""
    skills = list(dict.fromkeys(profile.get("skills", [])))
    tools = list(dict.fromkeys(profile.get("tools", [])))
    roles = list(dict.fromkeys(profile.get("roles", [])))

    experience = []
    for exp in profile.get("experience", []):
        experience.append(
            {
                "role": exp.get("role", ""),
                "company": exp.get("company", ""),
                "date": exp.get("date", "") or f"{exp.get('from', '')} - {exp.get('to', '')}".strip(" -"),
                "description": exp.get("description", ""),
            }
        )

    education = []
    for edu in profile.get("education", []):
        education.append(
            {
                "degree": edu.get("degree", ""),
                "field": edu.get("field", ""),
                "college": edu.get("college", ""),
                "date": edu.get("date", "") or f"{edu.get('from', '')} - {edu.get('to', '')}".strip(" -"),
            }
        )

    return {
        "name": profile.get("name", ""),
        "dob": profile.get("dob", ""),
        "phone": profile.get("phone", ""),
        "email": profile.get("email", ""),
        "location": profile.get("location", ""),
        "linkedin": profile.get("linkedin", ""),
        "github": profile.get("github", ""),
        "summary": profile.get("summary", ""),
        "skills": skills,
        "tools": tools,
        "roles": roles,
        "experience": experience,
        "education": education,
        "projects": profile.get("projects", []),
    }


def _set_global_style(doc):
    # ATS-safe styling: simple font, no tables, no columns, no text boxes.
    style = doc.styles["Normal"]
    style.font.name = "Garamond"
    style.element.rPr.rFonts.set(qn("w:ascii"), "Garamond")
    style.element.rPr.rFonts.set(qn("w:hAnsi"), "Garamond")
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Garamond")
    style.element.rPr.rFonts.set(qn("w:cs"), "Garamond")
    style.font.size = Pt(11)


def _non_empty(values):
    return [v for v in values if str(v).strip()]


def _format_single_date(value):
    text = str(value or "").strip()
    if not text:
        return ""

    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%d/%m/%Y", "%Y/%m/%d", "%d %b %Y", "%d %B %Y", "%b %Y", "%B %Y"):
        try:
            parsed = datetime.strptime(text, fmt)
            if fmt in ("%b %Y", "%B %Y"):
                return parsed.strftime("%b %Y")
            return parsed.strftime("%d %b %Y")
        except ValueError:
            continue

    return text


def _format_date_value(value):
    text = str(value or "").strip()
    if not text:
        return ""

    parts = re.split(r"\s*(?:-|–|—|to)\s*", text, maxsplit=1, flags=re.IGNORECASE)
    if len(parts) == 2:
        start, end = parts
        formatted_start = _format_single_date(start)
        formatted_end = _format_single_date(end)

        if formatted_start != start or formatted_end != end:
            return f"{formatted_start} - {formatted_end}"

    return _format_single_date(text)


def _split_bullets(text):
    if isinstance(text, list):
        raw = [str(x).strip() for x in text if str(x).strip()]
        return raw
    text = str(text or "").strip()
    if not text:
        return []

    if "\n" in text or "\r" in text or "•" in text:
        lines = re.split(r"[\n\r]+", text)
        bullets = []
        for line in lines:
            cleaned = re.sub(r"^[\-•\*\s]+", "", line).strip()
            if cleaned:
                bullets.append(cleaned)
        return bullets

    parts = re.split(r"\s*[;|]\s*|\s*\u2022\s*", text)
    bullets = [p.strip(" -") for p in parts if p.strip(" -")]
    if not bullets:
        return []
    if len(bullets) == 1:
        return bullets
    return bullets


def _add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True

    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_header(doc, data):
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = name.add_run(data.get("name", ""))
    run.bold = True
    run.font.size = Pt(16)

    contact_items = _non_empty(
        [
            data.get("email", ""),
            data.get("phone", ""),
            _format_date_value(data.get("dob", "")),
            data.get("location", ""),
            data.get("linkedin", ""),
            data.get("github", ""),
        ]
    )
    if contact_items:
        doc.add_paragraph(" | ".join(contact_items))


def _build_summary(data):
    top_skills = data.get("skills", [])[:6]
    roles = data.get("roles", [])[:2]
    exp_count = len(data.get("experience", []))

    role_text = ", ".join(roles) if roles else "software and AI roles"
    skill_text = ", ".join(top_skills) if top_skills else "problem solving, collaboration, and technical execution"
    exp_text = f"with {exp_count} practical experience entr{'y' if exp_count == 1 else 'ies'}" if exp_count else "with project-focused hands-on training"

    return (
        f"Results-driven candidate targeting {role_text}, {exp_text}; "
        f"strong in {skill_text}; focused on building production-ready, scalable solutions."
    )


def _add_summary(doc, data):
    _add_heading(doc, "Professional Summary")
    custom_summary = str(data.get("summary", "")).strip()
    doc.add_paragraph(custom_summary if custom_summary else _build_summary(data))


def _add_skills(doc, data):
    _add_heading(doc, "Technical Skills")

    skills = data.get("skills", [])
    tools = data.get("tools", [])
    roles = data.get("roles", [])

    if skills:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Core Skills: ").bold = True
        p.add_run(", ".join(skills))

    if tools:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Tools/Platforms: ").bold = True
        p.add_run(", ".join(tools))

    if roles:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("Target Roles: ").bold = True
        p.add_run(", ".join(roles))


def _add_experience(doc, data):
    experience = data.get("experience", [])
    if not experience:
        return

    _add_heading(doc, "Experience")
    for idx, exp in enumerate(experience):
        role = exp.get("role", "")
        company = exp.get("company", "")
        date = _format_date_value(exp.get("date", ""))

        title = " | ".join(_non_empty([role, company]))
        if date:
            title = f"{title} | {date}" if title else date

        if title:
            t = doc.add_paragraph()
            t.add_run(title).bold = True

        bullets = _split_bullets(exp.get("description", ""))
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

        if idx < len(experience) - 1:
            doc.add_paragraph()


def _add_projects(doc, data):
    projects = data.get("projects", [])
    if not projects:
        return

    _add_heading(doc, "Projects")
    for proj in projects:
        title = proj.get("name", "")
        if title:
            p = doc.add_paragraph()
            p.add_run(title).bold = True

        bullets = _split_bullets(proj.get("description", []))
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")


def _add_education(doc, data):
    _add_heading(doc, "Education")
    for edu in data.get("education", []):
        degree = edu.get("degree", "")
        field = edu.get("field", "")
        college = edu.get("college", "")
        date = _format_date_value(edu.get("date", ""))

        degree_line = degree
        if field:
            degree_line = f"{degree} in {field}" if degree else field

        line_one_parts = _non_empty([degree_line, college, date])
        if line_one_parts:
            p = doc.add_paragraph()
            if degree_line:
                p.add_run(degree_line).bold = True
            remaining_parts = _non_empty([college, date])
            if remaining_parts:
                if degree_line:
                    p.add_run(" | ")
                p.add_run(" | ".join(remaining_parts))


def generate_resume(data, output_file="ats_resume.docx"):
    """Generate an ATS-friendly resume DOCX from structured profile data."""
    doc = Document()
    _set_global_style(doc)

    _add_header(doc, data)
    _add_summary(doc, data)
    _add_skills(doc, data)
    _add_experience(doc, data)
    _add_projects(doc, data)
    _add_education(doc, data)

    doc.save(output_file)
    print(f"ATS resume generated: {output_file}")


if __name__ == "__main__":
    sample = {
        "name": "Candidate Name",
        "phone": "+91-9000000000",
        "email": "candidate@email.com",
        "location": "Bengaluru, India",
        "linkedin": "linkedin.com/in/candidate",
        "github": "github.com/candidate",
        "skills": ["Python", "Machine Learning", "FastAPI", "Docker"],
        "tools": ["Git", "Linux", "PyTorch"],
        "roles": ["Machine Learning Engineer"],
        "experience": [
            {
                "role": "AI Intern",
                "company": "Example Pvt Ltd",
                "date": "May 2025 - Jul 2025",
                "description": "Built model APIs; optimized inference for edge deployment; documented architecture.",
            }
        ],
        "education": [
            {
                "degree": "B.Tech",
                "field": "Computer Science",
                "college": "VIT",
                "date": "2022 - 2026",
            }
        ],
        "projects": [],
    }
    generate_resume(sample, "ats_resume.docx")